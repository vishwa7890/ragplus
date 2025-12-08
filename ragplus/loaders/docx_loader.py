"""
DOCX document loader using python-docx (fully offline).
"""

try:
    from docx import Document
except ImportError:
    Document = None


def load_docx(file_path: str) -> str:
    """
    Load text from a DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If file doesn't exist
    """
    if Document is None:
        raise ImportError("Install python-docx: pip install python-docx")
    
    doc = Document(file_path)
    
    # Extract text from paragraphs
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    
    # Extract text from tables
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_text))
    
    # Combine all text
    all_text = "\n".join(paragraphs)
    if table_text:
        all_text += "\n\nTables:\n" + "\n".join(table_text)
    
    return all_text.strip()
